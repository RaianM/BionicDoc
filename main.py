import tkinter as tk
from tkinter import filedialog, Tk, font
from docx import Document
from docx.shared import Pt
from hyphen import Hyphenator

def split_into_syllables(word):
    # Syllable splitting using the Hunspell algorithm via pyhyphen
    hyphenator = Hyphenator('en_US')
    syllables = hyphenator.syllables(word) # FIX SYLLABLE STUFF AND MODIFY
    return syllables

def process_file(file_path, output_name):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        new_runs = []
        for run in paragraph.runs:
            words = run.text.split()

            for i, word in enumerate(words):
                syllables = split_into_syllables(word)
                print(syllables)

                if(len(syllables) <= 1):
                    #find midpoint, bold left side
                    midpoint = len(word) // 2
                    bold_run = paragraph.add_run(word[:midpoint])
                    bold_run.font.bold = True
                    new_runs.append(bold_run)
                    unbold_run = paragraph.add_run(word[midpoint:])
                    new_runs.append(unbold_run)
                else:
                    for j, syllable in enumerate(syllables):
                        syllable_run = paragraph.add_run(syllable)
                        if j == 0:  # Bold the first syllable
                            syllable_run.font.bold = True
                        new_runs.append(syllable_run)

                if i < len(words) - 1:  # Add a space between words
                    space_run = paragraph.add_run(" ")
                    new_runs.append(space_run)

            run.clear()

        for new_run in new_runs:
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.italic = run.font.italic
            new_run.font.underline = run.font.underline

    output_path = f"{output_name}.docx"
    doc.save(output_path)
    print("File processed successfully.")

def select_file():
    file_path = filedialog.askopenfilename()
    output_name = output_text.get()  # Get the entered output name from the entry field
    process_file(file_path, output_name)

def quit_program():
    print("Quitting the program...")
    window.quit()

def main():
    bg_color = "#333e7a"
    fg_color = "#f2f2f2"

    global window
    window = tk.Tk()
    window.title("BionicDoc")  # Set window title
    window.geometry("800x600")  # Set window size (width x height)
    window.configure(background = bg_color)
    window.resizable(width=True, height=True)

    custom_font = "Lucida Sans Unicode"

    # Create the menu options
    label = tk.Label(window, text="BionicDoc", font=(custom_font, 20, "bold"), bg=bg_color, fg=fg_color)  # Increase font size and add padding
    label.pack(pady=(50,20))

    text_label = tk.Label(window, text="Insert the document you would like to modify into the Bionic Reading Style.\nThe modified document will be placed in the same folder as the original. ", font=(custom_font, 12), padx=50, bg=bg_color, fg=fg_color)  # Increase font size and add padding
    text_label.pack(pady=(0, 15))

    entry_label = tk.Label(window, text="New Document's Name:", font=(custom_font, 14), bg=bg_color, fg=fg_color)
    entry_label.pack()

    global output_text 
    output_text  = tk.Entry(window, font=(("Courier New", 20, "bold"), 12), highlightthickness=1, highlightbackground=fg_color, width=20, bg=fg_color, fg=bg_color)  # Increase the width of the entry field
    output_text.pack()

    button_frame = tk.Frame(window, bg=bg_color)  # Create a frame for the buttons
    button_frame.pack(pady=20)

    button1 = tk.Button(button_frame, text="Input a file", font=("Arial", 14), padx=10, pady=5, relief=tk.RAISED, bd=2, bg=bg_color, fg=fg_color, command=select_file)  # Add padding, relief, and border width
    button1.pack(side=tk.LEFT, padx=10)

    button2 = tk.Button(button_frame, text="Quit", font=("Arial", 14), padx=10, pady=5, relief=tk.RAISED, bd=2, bg=bg_color, fg=fg_color, command=quit_program)
    button2.pack(side=tk.LEFT, padx=10)

    # padx/pady = padding, relief style = styling around button, bd = border width, fg = text color

    # Start the main event loop
    window.mainloop()

if __name__ == "__main__":
    main()

# #3366cc - BLUE, #800000 - MAROON, "#85a5d4" - LIGHT BLUE, #E6E3D3 - CREAM, 

"""doc = Document(file_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            words = run.text.split()

            odd = True
            new_runs = []
            for word in words:
                if(odd):
                    midpoint = len(word) // 2
                    bold_run = paragraph.add_run(word[:midpoint])
                    bold_run.font.bold = True
                    new_runs.append(bold_run)

                    unbold_run = paragraph.add_run(word[midpoint:])
                    new_runs.append(unbold_run)
                    
                    #space_run = paragraph.add_run(" ")
                    #new_runs.append(space_run)
                    odd = False
                else:
                    new_runs.append(paragraph.add_run(word))
                    odd = True 

                space_run = paragraph.add_run(" ")
                new_runs.append(space_run)

            for i, word_run in enumerate(new_runs):
                if i % 2 == 0:  # Set font properties for the bolded runs
                    word_run.font.name = run.font.name
                    word_run.font.size = run.font.size
                    word_run.font.color.rgb = run.font.color.rgb
                    word_run.font.italic = run.font.italic
                    word_run.font.underline = run.font.underline
                else:  # Set font properties for the unbolded runs
                    word_run.font.name = run.font.name
                    word_run.font.size = run.font.size
                    word_run.font.color.rgb = run.font.color.rgb

            run.clear() """